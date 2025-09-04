#!/usr/bin/env python3
"""
PDF to Word Converter with Editable Text
Standalone Python script for ConvertWiz PDF to Word conversion
Uses PyMuPDF for text extraction and python-docx for Word creation with editable content
"""

import sys
import os
import pymupdf as fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import re

def extract_structured_text_from_page(page):
    """Extract text with structure detection for Word organization"""
    try:
        # Get text blocks with positioning information
        blocks = page.get_text("dict")
        structured_data = []
        
        for block in blocks["blocks"]:
            if "lines" in block:
                for line in block["lines"]:
                    line_text = ""
                    font_size = 12
                    is_bold = False
                    
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if text:
                            line_text += text + " "
                            font_size = max(font_size, span["size"])
                            if span["flags"] & 2**4:  # Bold flag
                                is_bold = True
                    
                    if line_text.strip():
                        structured_data.append({
                            "text": line_text.strip(),
                            "font_size": font_size,
                            "is_bold": is_bold
                        })
        
        return structured_data
    except:
        # Fallback: simple text extraction
        text = page.get_text()
        if text.strip():
            lines = text.split('\n')
            structured_data = []
            for line in lines:
                if line.strip():
                    structured_data.append({
                        "text": line.strip(),
                        "font_size": 12,
                        "is_bold": False
                    })
            return structured_data
        
        return []

def detect_headers(structured_data):
    """Detect potential headers based on font size and content"""
    if not structured_data:
        return []
    
    # Calculate average font size
    avg_font_size = sum(item["font_size"] for item in structured_data) / len(structured_data)
    
    headers = []
    for item in structured_data:
        # Potential header if significantly larger than average or bold
        if item["font_size"] > avg_font_size + 2 or item["is_bold"]:
            # Also check if it's short (likely a header)
            if len(item["text"]) < 100:
                headers.append(item)
    
    return headers

def add_text_to_word(document, structured_data, page_num):
    """Add structured text to Word document with formatting"""
    if not structured_data:
        # Add placeholder for empty pages
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Page {page_num}: No extractable text (may be image-based)")
        run.font.italic = True
        run.font.color.rgb = RGBColor(102, 102, 102)
        return
    
    # Add page separator for multi-page documents
    if page_num > 1:
        document.add_page_break()
        header_para = document.add_paragraph()
        header_run = header_para.add_run(f"--- Page {page_num} ---")
        header_run.font.bold = True
        header_run.font.size = Pt(14)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph()  # Empty line after header
    
    # Detect headers for better formatting
    headers = detect_headers(structured_data)
    header_texts = {item["text"] for item in headers}
    
    for item in structured_data:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(item["text"])
        
        # Apply formatting based on content analysis
        if item["text"] in header_texts:
            # Format as header
            run.font.bold = True
            run.font.size = Pt(max(14, min(item["font_size"], 18)))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Format as normal text
            if item["is_bold"]:
                run.font.bold = True
            run.font.size = Pt(max(10, min(item["font_size"], 14)))
        
        # Add some spacing for readability
        paragraph.space_after = Pt(6)

def convert_pdf_to_word(input_path, output_path):
    """Convert PDF to Word document with editable text"""
    try:
        print(f"Starting conversion: {input_path} -> {output_path}")
        start_time = time.time()
        
        # Open PDF document
        pdf_document = fitz.open(input_path)
        
        # Create new Word document
        document = Document()
        
        # Add title
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run("Converted from PDF")
        title_run.font.bold = True
        title_run.font.size = Pt(16)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document info
        info_paragraph = document.add_paragraph()
        info_run = info_paragraph.add_run(f"Source: {os.path.basename(input_path)}")
        info_run.font.italic = True
        info_run.font.size = Pt(10)
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator
        document.add_paragraph("=" * 50)
        
        # Process each page of the PDF
        total_pages = len(pdf_document)
        for page_num in range(total_pages):
            print(f"Processing page {page_num + 1}/{total_pages}")
            
            # Get the page
            page = pdf_document.load_page(page_num)
            
            # Extract structured text from the page
            structured_data = extract_structured_text_from_page(page)
            
            # Add text to Word document
            add_text_to_word(document, structured_data, page_num + 1)
        
        # Close PDF document
        pdf_document.close()
        
        # Add footer with conversion info
        footer_paragraph = document.add_paragraph()
        footer_paragraph.add_run("=" * 50)
        
        summary_paragraph = document.add_paragraph()
        summary_run = summary_paragraph.add_run(f"Conversion completed: {total_pages} pages processed")
        summary_run.font.bold = True
        summary_run.font.size = Pt(10)
        summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save Word document
        document.save(output_path)
        
        duration = round(time.time() - start_time, 2)
        print(f"Conversion completed in {duration} seconds")
        
        # Check if output file exists
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"Output file created: {output_path} ({file_size} bytes)")
            return True
        else:
            print("Error: Output file not created")
            return False
            
    except Exception as e:
        print(f"Conversion error: {str(e)}")
        return False

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python pdf_to_word_converter.py <input_pdf> <output_docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)
    
    success = convert_pdf_to_word(input_file, output_file)
    sys.exit(0 if success else 1)